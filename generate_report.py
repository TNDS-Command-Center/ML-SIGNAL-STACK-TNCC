"""
generate_report.py — SignalStack: Automated Word Document Report Generator.
============================================================================
Reads live metrics.txt and forecast_results.csv from each source's output
folder, then builds a fully formatted .docx intelligence report with:
  - Executive summary table
  - Per-source forecast charts (embedded PNG)
  - Model accuracy scores with plain-English interpretation
  - Glossary

Usage:
    python generate_report.py
    python generate_report.py --out "reports/SignalStack_Report_2025-W16.docx"
    python generate_report.py --source sales ops_pulse

After any pipeline run, call this to get a fresh report in seconds.
"""

import os
import sys
import argparse
import re
import datetime
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── Config ───────────────────────────────────────────────────────────────────

BASE_DIR = Path(__file__).parent

NAVY  = RGBColor(0x1A, 0x3A, 0x5C)
TEAL  = RGBColor(0x3D, 0x8E, 0xB9)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
LGRAY = RGBColor(0xF5, 0xF5, 0xF5)
BLACK = RGBColor(0x00, 0x00, 0x00)
GREEN = RGBColor(0x3B, 0x6D, 0x11)
AMBER = RGBColor(0x85, 0x4F, 0x0B)
RED   = RGBColor(0xA3, 0x2D, 0x2D)

SOURCES = ["sales", "ops_pulse", "cash_flow_compass", "pipeline_pulse", "team_tempo"]

SOURCE_META = {
    "sales":             {"label": "Daily Sales Revenue",          "unit": "$",    "target": "Total Sales"},
    "ops_pulse":         {"label": "Weekly Jobs Completed",        "unit": "jobs", "target": "Jobs Done"},
    "cash_flow_compass": {"label": "Weekly Ending Cash Balance",   "unit": "$",    "target": "Ending Balance"},
    "pipeline_pulse":    {"label": "Weekly Pipeline Value",        "unit": "$",    "target": "Pipeline Value"},
    "team_tempo":        {"label": "Weekly Billable Hours",        "unit": "hrs",  "target": "Billable Hrs"},
}

ACCURACY_BANDS = [
    (5,   "Excellent — high-confidence forecast. Use for operational decisions."),
    (10,  "Good — reliable signal. Minor variance expected week-to-week."),
    (20,  "Moderate — directionally accurate. Verify against actuals each week."),
    (35,  "Fair — use for trend direction only. More data will improve accuracy."),
    (100, "Developing — baseline established. Model improves as history builds."),
]

GLOSSARY = [
    ("SARIMA",             "Seasonal AutoRegressive Integrated Moving Average — the statistical model used to learn patterns in time series data and project them forward."),
    ("AIC",                "Akaike Information Criterion — a model quality score. Lower AIC = better fit. Used to select the best parameter combination during grid search."),
    ("MAE (Mean Absolute Error)", "Average dollar/unit difference between forecast and actual. Easy to interpret: if MAE = $307, the forecast is off by $307 on average."),
    ("RMSE (Root Mean Square Error)", "Similar to MAE but penalizes large misses more heavily. Higher than MAE means there are occasional big misses."),
    ("MAPE (%)",           "Mean Absolute Percentage Error — forecast accuracy expressed as a percentage. 5% MAPE means the model is off by 5% on average. Under 10% is considered good for business forecasting."),
    ("Confidence Interval (95% CI)", "The shaded band on forecast charts. There is a 95% probability the actual value will land inside this band. Wider band = more uncertainty."),
    ("Training Data",      "The historical period the model learned from (blue line on charts)."),
    ("Validation Period",  "The held-out period used to test model accuracy before forecasting the future (green line on charts — this is real data the model had not seen)."),
    ("Forecast",           "The model's projection beyond the validation period (orange line on charts)."),
    ("Grid Search",        "The automated process of testing all SARIMA parameter combinations and selecting the one with the lowest AIC. Run once; reuse with --skip-search."),
    ("Ensemble Forecast",  "A blending technique that combines SARIMA with a Weighted Moving Average (WMA) to stabilize volatile signals. The SARIMA component captures seasonality and trend; the WMA anchors the forecast to recent observed levels. Used for pipeline value and similar high-variance signals."),
    ("Signal Profile",     "Business context about the data behind a forecast — customer count, transaction volume, revenue concentration. Explains why some signals are inherently harder to forecast than others."),
    ("ISO Week",           "A standardized week numbering system (e.g., 2025-W14 = the 14th week of 2025, starting Monday). Used for all weekly SignalStack sources."),
    ("EWM Smoothing",      "Exponential Weighted Moving Average — applied during preprocessing to reduce day-to-day noise while preserving trend direction."),
]


def load_glossary_from_manual(manual_path):
    """
    Load the full glossary section from docs/SIGNALSTACK_USER_MANUAL.md.

    Expected format in the glossary section:
        **Term**
        Definition paragraph...
    """
    path = Path(manual_path)
    if not path.exists():
        return []

    text = path.read_text(encoding="utf-8")
    start_match = re.search(r"^##\s+13\.\s+Glossary of Terms\s*$", text, flags=re.MULTILINE)
    if not start_match:
        return []

    # End section at the next top-level heading or markdown rule, whichever appears first.
    tail = text[start_match.end():]
    next_heading = re.search(r"^##\s+\d+\.\s+", tail, flags=re.MULTILINE)
    next_rule = re.search(r"^\s*---\s*$", tail, flags=re.MULTILINE)

    cut_positions = [len(tail)]
    if next_heading:
        cut_positions.append(next_heading.start())
    if next_rule:
        cut_positions.append(next_rule.start())
    section = tail[:min(cut_positions)]

    entries = []
    lines = section.splitlines()
    idx = 0
    while idx < len(lines):
        line = lines[idx].strip()
        term_match = re.match(r"^\*\*(.+?)\*\*$", line)
        if not term_match:
            idx += 1
            continue

        term = term_match.group(1).strip()
        idx += 1
        definition_lines = []
        while idx < len(lines):
            candidate = lines[idx].strip()
            if re.match(r"^\*\*(.+?)\*\*$", candidate):
                break
            if candidate:
                definition_lines.append(candidate)
            idx += 1

        definition = " ".join(definition_lines).strip()
        if definition:
            entries.append((term, definition))

    return entries


# Prefer full glossary from the user manual for DOCX + HTML outputs.
_manual_glossary = load_glossary_from_manual(BASE_DIR / "docs" / "SIGNALSTACK_USER_MANUAL.md")
if _manual_glossary:
    # Keep full manual glossary first, then append any report-specific terms not present.
    known_terms = {term for term, _ in _manual_glossary}
    merged_glossary = list(_manual_glossary)
    for term, definition in GLOSSARY:
        if term not in known_terms:
            merged_glossary.append((term, definition))
    GLOSSARY = merged_glossary

# ── Helpers ───────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)

def cell_para(cell, text, bold=False, size=10, color=BLACK, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.color.rgb = color
    return p

def heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.style = f"Heading {level}"
    run = p.add_run(text)
    if level == 1:
        run.font.color.rgb = NAVY
        run.font.size = Pt(16)
        run.bold = True
    elif level == 2:
        run.font.color.rgb = TEAL
        run.font.size = Pt(13)
        run.bold = True
    return p

def body(doc, text, size=10.5, color=BLACK, bold=False, italic=False, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic
    return p

def rule(doc, color_hex="3D8EB9"):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(8)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def accuracy_label(mape):
    for threshold, label in ACCURACY_BANDS:
        if mape <= threshold:
            return label
    return ACCURACY_BANDS[-1][1]

def accuracy_color(mape):
    if mape <= 10:
        return GREEN
    if mape <= 20:
        return AMBER
    return RED

FREQ_UNITS = {"B": "business days", "W": "weeks", "D": "days", "M": "months"}
CYCLE_NAMES = {
    ("B", 5): "one trading week", ("W", 4): "roughly one month",
    ("W", 12): "one quarter", ("W", 52): "one year",
    ("M", 12): "one year", ("D", 7): "one week",
}


def parse_sarima(model_str):
    """Parse SARIMA(p,d,q)x(P,D,Q,s) string into component ints."""
    match = re.match(
        r"SARIMA\((\d+),\s*(\d+),\s*(\d+)\)x\((\d+),\s*(\d+),\s*(\d+),\s*(\d+)\)",
        model_str,
    )
    if not match:
        return None
    return tuple(int(x) for x in match.groups())


def narrate_model(model_str, freq):
    """Paragraph 2 — plain-English model explanation."""
    parsed = parse_sarima(model_str)
    if not parsed:
        return f"Model specification: {model_str}."
    p, d, q, P, D, Q, s = parsed
    sentences = []

    # Differencing
    if d == 0:
        sentences.append("No differencing was needed — the data was already stationary enough for modeling.")
    elif d == 1:
        sentences.append("The model applies one differencing pass to remove the underlying trend slope before learning patterns.")
    else:
        sentences.append(f"The model applies {d} differencing passes to remove an accelerating trend before learning patterns.")

    # Seasonal cycle
    fu = FREQ_UNITS.get(freq, "periods")
    cycle = CYCLE_NAMES.get((freq, s))
    if cycle:
        sentences.append(f"It looks for repeating cycles every {s} {fu}, which corresponds to {cycle} in this data.")
    else:
        sentences.append(f"It looks for repeating cycles every {s} {fu} in this data.")

    # AR/MA components
    if p == 0 and q == 0 and P == 0 and Q == 0:
        sentences.append(
            "All autoregressive and moving-average components are zeroed out — the model "
            "found only directional momentum in the training window and will improve as "
            "history accumulates."
        )
    else:
        active = []
        if p > 0:
            active.append("short-term momentum")
        if q > 0:
            active.append("recent forecast-error correction")
        if P > 0:
            active.append("seasonal momentum")
        if Q > 0:
            active.append("seasonal error correction")
        sentences.append(f"Active components: {', '.join(active)}.")

    return " ".join(sentences)


def narrate_ensemble(source, model_str, freq):
    """
    If the source uses ensemble forecasting, prepend an ensemble explanation
    to the standard model narration.
    """
    from config import get_source
    try:
        src = get_source(source)
    except KeyError:
        return narrate_model(model_str, freq)

    if not src.get("ensemble_forecast"):
        return narrate_model(model_str, freq)

    w_sarima, w_ma = src.get("ensemble_weights", (0.6, 0.4))
    base = narrate_model(model_str, freq)
    ensemble_note = (
        f"This signal uses an ensemble approach: the SARIMA forecast is blended "
        f"with a weighted moving average at a {int(w_sarima*100)}/{int(w_ma*100)} "
        f"ratio. The WMA component anchors the projection to recent observed levels, "
        f"reducing the impact of week-to-week volatility on the forecast line."
    )
    return f"{ensemble_note} {base}"


def narrate_numbers(m, meta):
    """Paragraph 3 — MAE / RMSE / MAPE in business language."""
    mape = float(m.get("MAPE", 0))
    mae = float(m.get("MAE", 0))
    rmse = float(m.get("RMSE", 0))
    avg = float(m.get("Average_Actual", 0))
    mae_pct = float(m.get("MAE_pct_of_avg", 0))
    unit = meta["unit"]

    if unit == "$":
        mae_fmt, rmse_fmt, avg_fmt = f"${mae:,.0f}", f"${rmse:,.0f}", f"${avg:,.0f}"
    else:
        mae_fmt = f"{mae:.1f} {unit}"
        rmse_fmt = f"{rmse:.1f} {unit}"
        avg_fmt = f"{avg:.1f} {unit}"

    sentences = []
    sentences.append(
        f"The average forecast error is {mae_fmt}, which represents {mae_pct:.1f}% "
        f"of the observed average ({avg_fmt})."
    )

    if rmse > mae * 1.3:
        sentences.append(
            f"RMSE ({rmse_fmt}) is notably higher than MAE, indicating occasional "
            f"large misses alongside the more typical errors."
        )
    else:
        sentences.append(
            f"RMSE ({rmse_fmt}) is close to MAE, meaning errors are consistent in "
            f"size without large outlier misses."
        )

    if mape <= 5:
        sentences.append(f"At {mape:.1f}% MAPE, this forecast is highly precise — expect actual values to land very close to projections.")
    elif mape <= 10:
        sentences.append(f"At {mape:.1f}% MAPE, this forecast is reliable for week-to-week planning with only minor variance expected.")
    elif mape <= 20:
        sentences.append(f"At {mape:.1f}% MAPE, the model captures the right direction but individual period values may vary moderately.")
    elif mape <= 35:
        sentences.append(f"At {mape:.1f}% MAPE, the model is still learning this signal's patterns — treat specific values as directional guidance.")
    else:
        sentences.append(f"At {mape:.1f}% MAPE, the model is in early development — forecast direction may be useful but magnitude is unreliable.")

    cv_mape = float(m.get("CV_Mean_MAPE", 0))
    cv_std = float(m.get("CV_Std_MAPE", 0))
    if cv_mape > 0 and cv_mape != float("inf"):
        sentences.append(
            f"Cross-validation ({3}-fold) confirms this at {cv_mape:.1f}% +/- {cv_std:.1f}% average MAPE, "
            f"indicating {'stable' if cv_std < 5 else 'variable'} model performance across different time windows."
        )

    bias_raw = m.get("Bias_Detected", False)
    if isinstance(bias_raw, str):
        bias_detected = bias_raw.strip().lower() in {"true", "1", "yes"}
    else:
        bias_detected = bool(bias_raw)
    bias_pattern = str(m.get("Bias_Pattern", "none"))
    if bias_detected and bias_pattern != "none":
        sentences.append(
            f"Note: A {bias_pattern} residual bias was detected during validation, "
            f"indicating the model's trend-tracking could be improved with additional "
            f"training data or parameter adjustment."
        )

    return " ".join(sentences)


def narrate_chart(label):
    """Paragraph 4 — what to look for in the forecast chart."""
    return (
        f"In the forecast chart below for {label}, the blue line shows the training "
        f"period the model learned from, the green line shows real validation data the "
        f"model had not seen (the honest test), and the orange line is the model's "
        f"forecast. If green tracks orange closely the model is reliable for that horizon; "
        f"if they diverge significantly, the business shifted in a way the model did not "
        f"anticipate. The shaded band represents the 95% confidence interval — a wider "
        f"band signals greater uncertainty."
    )


def narrate_decision(mape, unit, mae):
    """Paragraph 5 — action sentence keyed to MAPE band."""
    if unit == "$":
        mae_fmt = f"${mae:,.0f}"
    else:
        mae_fmt = f"{mae:.1f} {unit}"

    if mape <= 10:
        return (
            f"Use this forecast for operational planning. At {mape:.1f}% MAPE it is "
            f"reliable enough to inform staffing, scheduling, and cash decisions directly."
        )
    elif mape <= 20:
        return (
            f"Use for directional awareness, not precise planning. Verify actuals weekly. "
            f"Flag if actual values diverge from the forecast by more than {mae_fmt} for "
            f"two consecutive periods."
        )
    elif mape <= 35:
        return (
            f"Treat as a trend indicator only — confirm whether the signal is rising, "
            f"flat, or falling. Do not use specific forecast values for budget or staffing "
            f"decisions until MAPE drops below 15%."
        )
    else:
        return (
            f"This model is still developing. Forecast direction may be useful but "
            f"individual values should not drive decisions. Priority: add more historical data."
        )


def narrate_signal_profile(source, meta):
    """
    Generate a Signal Profile paragraph for sources with an Analysis tab.
    Reads key business metrics from the Analysis sheet to contextualize
    model accuracy in business terms.

    Only applies to 'sales' for now. Returns empty string for other sources.
    """
    if source != "sales":
        return ""

    try:
        wb_path = BASE_DIR / "tnds-sales-data-template.xlsx"
        if not wb_path.exists():
            return ""

        import openpyxl
        wb = openpyxl.load_workbook(str(wb_path), data_only=True)
        if "Analysis" not in wb.sheetnames:
            return ""
        ws = wb["Analysis"]

        total_rev = ws["B6"].value
        transactions = ws["B7"].value
        unique_customers = ws["H8"].value
        avg_sale = ws["E6"].value
        top5_share = ws["B53"].value  # Top 5 Revenue Share (decimal)

        if not all([total_rev, transactions, unique_customers]):
            return ""

        sentences = []
        sentences.append(
            f"This forecast is built from {int(transactions):,} historical transactions "
            f"across {int(unique_customers)} unique customers, with total tracked revenue "
            f"of ${total_rev:,.0f}."
        )

        if top5_share and top5_share > 0:
            sentences.append(
                f"The top 5 customers account for {top5_share:.0%} of all revenue. "
                f"This concentration means individual ordering patterns directly impact "
                f"daily revenue variance — a structural data characteristic, not a model "
                f"weakness. Accuracy will improve as customer count grows."
            )

        if avg_sale and avg_sale > 500:
            sentences.append(
                f"Average transaction size is ${avg_sale:,.0f}, indicating bulk/commercial "
                f"purchasing patterns typical of field service operations."
            )

        return " ".join(sentences)
    except Exception:
        return ""


def load_metrics(source):
    path = BASE_DIR / "data" / "output" / source / "metrics.txt"
    if not path.exists():
        return None
    m = {}
    for line in path.read_text().splitlines():
        if ": " in line:
            k, v = line.split(": ", 1)
            try:
                m[k.strip()] = float(v.strip())
            except ValueError:
                m[k.strip()] = v.strip()
    return m

# ── Report builder ────────────────────────────────────────────────────────────

def build_report(sources, out_path):
    doc = Document()

    # Page setup — US Letter, 1" margins
    section = doc.sections[0]
    section.page_width  = Emu(12240 * 914)  # 8.5"
    section.page_height = Emu(15840 * 914)  # 11"
    section.left_margin   = Inches(1)
    section.right_margin  = Inches(1)
    section.top_margin    = Inches(1)
    section.bottom_margin = Inches(0.75)

    # Default font
    doc.styles["Normal"].font.name = "Arial"
    doc.styles["Normal"].font.size = Pt(10.5)

    # ── Cover block ───────────────────────────────────────────────────────────
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run("SIGNALSTACK")
    run.font.name = "Arial"
    run.font.size = Pt(28)
    run.font.bold = True
    run.font.color.rgb = NAVY

    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(2)
    run2 = p2.add_run("Business Intelligence Forecast Report")
    run2.font.size = Pt(14)
    run2.font.color.rgb = TEAL
    run2.font.name = "Arial"

    today = datetime.date.today()
    iso_week = today.strftime("%G-W%V")
    p3 = doc.add_paragraph()
    p3.paragraph_format.space_after = Pt(20)
    run3 = p3.add_run(f"Generated: {today.strftime('%B %d, %Y')}  |  Week: {iso_week}  |  True North Data Strategies")
    run3.font.size = Pt(9)
    run3.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    run3.font.name = "Arial"

    rule(doc, "1A3A5C")

    # ── Executive summary ─────────────────────────────────────────────────────
    heading(doc, "Executive Summary", 1)
    body(doc, "This report summarizes the latest SignalStack forecasting results across all active business signals. "
              "Each model was trained on historical data and validated against a held-out period before generating "
              "forward projections. Use MAPE as the primary accuracy indicator — under 10% is considered strong "
              "for small-business operational data.", space_after=10)

    # Summary table
    tbl = doc.add_table(rows=1, cols=5)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    col_widths_dxa = [2800, 1600, 1400, 1400, 2200]
    col_widths_in  = [w / 1440 for w in col_widths_dxa]

    headers = ["Signal", "Model", "MAPE", "MAE", "Accuracy Assessment"]
    hdr_row = tbl.rows[0]
    for i, (cell, hdr) in enumerate(zip(hdr_row.cells, headers)):
        set_cell_bg(cell, "1A3A5C")
        cell.width = Inches(col_widths_in[i])
        cell_para(cell, hdr, bold=True, size=9, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)

    all_metrics = {}
    for source in sources:
        m = load_metrics(source)
        if not m:
            continue
        all_metrics[source] = m
        meta  = SOURCE_META[source]
        mape  = float(m.get("MAPE", 0))
        mae   = float(m.get("MAE", 0))
        model = str(m.get("Model", ""))
        acol  = accuracy_color(mape)
        alabel = accuracy_label(mape).split(" — ")[0]

        row = tbl.add_row()
        for i, cell in enumerate(row.cells):
            cell.width = Inches(col_widths_in[i])
        set_cell_bg(row.cells[0], "F0F6FA")
        cell_para(row.cells[0], meta["label"], bold=True, size=9)
        cell_para(row.cells[1], model.replace("SARIMA", "").strip(), size=8.5, align=WD_ALIGN_PARAGRAPH.CENTER)
        cell_para(row.cells[2], f"{mape:.1f}%", bold=True, size=9,
                  color=acol, align=WD_ALIGN_PARAGRAPH.CENTER)
        unit = meta["unit"]
        mae_str = f"${mae:,.0f}" if unit == "$" else f"{mae:.1f} {unit}"
        cell_para(row.cells[3], mae_str, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
        cell_para(row.cells[4], alabel, size=9, color=acol)

    doc.add_paragraph()

    # ── Glossary (moved to top section) ──────────────────────────────────────
    heading(doc, "Glossary of Terms", 1)
    rule(doc, "1A3A5C")
    body(doc, "Reference definitions for all metrics, model parameters, and technical terms used in this report.",
         color=RGBColor(0x55, 0x55, 0x55), space_after=12)

    for term, definition in GLOSSARY:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.left_indent = Inches(0.2)
        term_run = p.add_run(f"{term}  ")
        term_run.bold = True
        term_run.font.color.rgb = NAVY
        term_run.font.size = Pt(10)
        def_run = p.add_run(definition)
        def_run.font.size = Pt(10)
        def_run.font.color.rgb = BLACK

    doc.add_paragraph()
    rule(doc, "3D8EB9")
    doc.add_page_break()

    # ── Per-source sections ───────────────────────────────────────────────────
    heading(doc, "Signal Details", 1)

    for source in sources:
        m = all_metrics.get(source)
        if not m:
            body(doc, f"[{source}] — no output data found. Run pipeline first.", italic=True)
            continue

        meta  = SOURCE_META[source]
        mape  = float(m.get("MAPE", 0))
        mae   = float(m.get("MAE", 0))
        rmse  = float(m.get("RMSE", 0))
        avg   = float(m.get("Average_Actual", 0))
        model = str(m.get("Model", ""))
        aic   = str(m.get("AIC", ""))
        unit  = meta["unit"]

        heading(doc, meta["label"], 2)
        rule(doc, "3D8EB9")

        # Metrics row table
        stats_tbl = doc.add_table(rows=2, cols=4)
        stats_tbl.style = "Table Grid"
        stat_labels = ["Model", "MAPE", "MAE", "RMSE"]
        mae_str  = f"${mae:,.0f}"  if unit == "$" else f"{mae:.1f}"
        rmse_str = f"${rmse:,.0f}" if unit == "$" else f"{rmse:.1f}"
        stat_vals   = [model, f"{mape:.1f}%", mae_str, rmse_str]
        stat_widths = [2.5, 1.3, 1.3, 1.3]

        for i, (lbl, val, w) in enumerate(zip(stat_labels, stat_vals, stat_widths)):
            lc = stats_tbl.rows[0].cells[i]
            vc = stats_tbl.rows[1].cells[i]
            lc.width = Inches(w)
            vc.width = Inches(w)
            set_cell_bg(lc, "1A3A5C")
            set_cell_bg(vc, "F0F6FA")
            cell_para(lc, lbl, bold=True, size=9, color=WHITE, align=WD_ALIGN_PARAGRAPH.CENTER)
            vcol = accuracy_color(mape) if lbl == "MAPE" else BLACK
            cell_para(vc, val, bold=(lbl == "MAPE"), size=10, color=vcol, align=WD_ALIGN_PARAGRAPH.CENTER)

        doc.add_paragraph()
        body(doc, accuracy_label(mape), italic=True, color=accuracy_color(mape), space_after=8)

        # ── Paragraph 2: Model explanation ──
        freq = str(m.get("Frequency", ""))
        body(doc, narrate_ensemble(source, model, freq), space_after=6)

        # ── Paragraph 3: Numbers in context ──
        body(doc, narrate_numbers(m, meta), space_after=6)

        # ── Paragraph 4: Chart narration ──
        body(doc, narrate_chart(meta["label"]), space_after=6)

        # ── Paragraph 5: Decision instruction ──
        decision_p = doc.add_paragraph()
        decision_p.paragraph_format.space_after = Pt(8)
        prefix_run = decision_p.add_run("Decision Guidance: ")
        prefix_run.bold = True
        prefix_run.font.size = Pt(10)
        prefix_run.font.color.rgb = NAVY
        decision_run = decision_p.add_run(narrate_decision(mape, unit, mae))
        decision_run.bold = True
        decision_run.font.size = Pt(10)
        decision_run.font.color.rgb = accuracy_color(mape)

        # ── Paragraph 6: Signal profile (sales only) ──
        profile = narrate_signal_profile(source, meta)
        if profile:
            profile_p = doc.add_paragraph()
            profile_p.paragraph_format.space_before = Pt(4)
            profile_p.paragraph_format.space_after = Pt(8)
            prefix_run = profile_p.add_run("Signal Profile: ")
            prefix_run.bold = True
            prefix_run.font.size = Pt(10)
            prefix_run.font.color.rgb = NAVY
            detail_run = profile_p.add_run(profile)
            detail_run.font.size = Pt(10)
            detail_run.font.color.rgb = BLACK

        # Embed forecast chart
        chart_path = BASE_DIR / "visuals" / source / "04_forecast_vs_actual.png"
        if chart_path.exists():
            try:
                doc.add_picture(str(chart_path), width=Inches(6.2))
                last_para = doc.paragraphs[-1]
                last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cap = doc.add_paragraph()
                cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cap.paragraph_format.space_after = Pt(14)
                from config import get_source
                src_cfg = get_source(source)
                chart_prefix = "Ensemble Forecast vs Actual" if src_cfg.get("ensemble_forecast") else "SARIMA Forecast vs Actual"
                cr = cap.add_run(f"Figure: {chart_prefix} — {meta['label']}")
                cr.font.size = Pt(8.5)
                cr.font.italic = True
                cr.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
            except Exception as e:
                body(doc, f"[Chart could not be embedded: {e}]", italic=True)
        else:
            body(doc, "[Chart not found — run pipeline to generate visuals]", italic=True)

        doc.add_paragraph()

    # ── Footer note ───────────────────────────────────────────────────────────
    footer_p = doc.add_paragraph()
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_p.paragraph_format.space_before = Pt(12)
    fr = footer_p.add_run(
        f"SignalStack by True North Data Strategies  |  jacob@truenorthstrategyops.com  |  719-204-6365  |  Colorado Springs, CO  |  SDVOSB"
    )
    fr.font.size = Pt(8)
    fr.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

    # ── Save ──────────────────────────────────────────────────────────────────
    out_path = Path(out_path)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    print(f"\n[report] Saved: {out_path}")
    print(f"[report] Sources included: {', '.join(sources)}")
    print(f"[report] Date: {today.strftime('%B %d, %Y')}  |  Week: {iso_week}\n")


# ── Entry point ───────────────────────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(
        description="SignalStack — Generate Word document report from latest pipeline output.",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
Examples:
  python generate_report.py
  python generate_report.py --out reports/SignalStack_2025-W16.docx
  python generate_report.py --source sales ops_pulse cash_flow_compass
        """
    )
    parser.add_argument(
        "--out",
        default=None,
        help="Output .docx path. Default: reports/SignalStack_Report_<YYYY-WNN>.docx"
    )
    parser.add_argument(
        "--source",
        nargs="+",
        default=SOURCES,
        choices=SOURCES,
        help="Sources to include. Default: all five."
    )
    args = parser.parse_args()

    today    = datetime.date.today()
    iso_week = today.strftime("%G-W%V")
    out_path = args.out or str(BASE_DIR / "reports" / f"SignalStack_Report_{iso_week}.docx")

    print(f"\n[report] Generating SignalStack report...")
    print(f"[report] Sources: {', '.join(args.source)}")
    print(f"[report] Output:  {out_path}")

    build_report(args.source, out_path)


if __name__ == "__main__":
    main()
